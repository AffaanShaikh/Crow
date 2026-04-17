"""
Multi-format document text extraction.

Supported formats:
  .txt, .md          - direct UTF-8 read
  .pdf               - pypdf (text layer; scanned PDFs need OCR, not supported here)
  .docx              - python-docx
  .html, .htm        - trafilatura (main content extraction, strips nav/ads)
  .py, .js, .ts etc. - treated as plain text (code is valid RAG content)

All parsers return a ParsedDocument containing the cleaned text and metadata.
Failed parses raise DocumentParseError w/ a message.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from utils.logger import get_logger 

log = get_logger(__name__)


class DocumentParseError(Exception):
    """Raised when a document cannot be parsed."""
    pass


@dataclass
class ParsedDocument:
    text: str
    title: str
    source_path: str
    file_type: str
    char_count: int = field(init=False)
    metadata: dict[str, Any] = field(default_factory=dict)

    def __post_init__(self):
        self.char_count = len(self.text)


def parse(path: Path | str, raw_bytes: bytes | None = None) -> ParsedDocument:
    """
    Parse a document into plain text.

    Args:
        path:      File path (used to determine format and as source reference).
        raw_bytes: If provided, parse from memory instead of reading from disk.
                   Useful for uploaded files before they are saved.

    Returns:
        ParsedDocument with extracted text and metadata.

    Raises:
        DocumentParseError: If parsing fails or the format is unsupported.
    """
    p = Path(path)
    suffix = p.suffix.lower()
    data = raw_bytes if raw_bytes is not None else p.read_bytes()

    try:
        if suffix in (".txt", ".md", ".rst", ".csv", ".json"):
            return _parse_text(data, p)
        elif suffix == ".pdf":
            return _parse_pdf(data, p)
        elif suffix in (".docx", ".doc"):
            return _parse_docx(data, p)
        elif suffix in (".html", ".htm"):
            return _parse_html(data, p)
        elif suffix in _CODE_EXTENSIONS:
            return _parse_text(data, p, file_type="code")
        else:
            # unknown extension - try as plain text
            return _parse_text(data, p, file_type="unknown")
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse '{p.name}': {exc}") from exc


# format-specific parsers

def _parse_text(data: bytes, path: Path, file_type: str = "text") -> ParsedDocument:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="replace")

    text = _clean_whitespace(text)
    if not text:
        raise DocumentParseError(f"'{path.name}' is empty after parsing")

    return ParsedDocument(
        text=text,
        title=path.stem,
        source_path=str(path),
        file_type=file_type,
    )


def _parse_pdf(data: bytes, path: Path) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise DocumentParseError(
            "pypdf not installed. Run: pip install pypdf"
        )

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []

    for page_num, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(page_text)
        except Exception:
            log.debug("pdf_page_extract_failed", filename=path.name, page=page_num)
            pass # skip unextractable pages silently

    if not pages:
        raise DocumentParseError(
            f"'{path.name}': no text could be extracted. "
            "The PDF may be scanned - OCR support is not available yet."
        )

    text = _clean_whitespace("\n\n".join(pages))
    return ParsedDocument(
        text=text,
        title=path.stem,
        source_path=str(path),
        file_type="pdf",
        metadata={"page_count": len(reader.pages), "extracted_pages": len(pages)},
    )


def _parse_docx(data: bytes, path: Path) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError:
        raise DocumentParseError(
            "python-docx not installed. Run: pip install python-docx"
        )

    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = _clean_whitespace("\n\n".join(paragraphs))

    if not text:
        raise DocumentParseError(f"'{path.name}' contains no text paragraphs")

    return ParsedDocument(
        text=text,
        title=path.stem,
        source_path=str(path),
        file_type="docx",
        metadata={"paragraph_count": len(paragraphs)},
    )


def _parse_html(data: bytes, path: Path) -> ParsedDocument:
    # Try trafilatura first (best for web content)
    try:
        import trafilatura
        text = trafilatura.extract(data.decode("utf-8", errors="replace"))
        if text and len(text.strip()) > 50:
            return ParsedDocument(
                text=_clean_whitespace(text),
                title=path.stem,
                source_path=str(path),
                file_type="html",
            )
    except ImportError:
        pass

    # fallback: beautifulsoup4
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(data, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        text = _clean_whitespace(text)
        if text:
            return ParsedDocument(
                text=text,
                title=path.stem,
                source_path=str(path),
                file_type="html",
            )
    except ImportError:
        pass

    # Last resort: strip HTML tags with regex
    import re
    text = re.sub(r"<[^>]+>", " ", data.decode("utf-8", errors="replace"))
    text = _clean_whitespace(text)
    if not text:
        raise DocumentParseError(f"'{path.name}': no text could be extracted from HTML")

    return ParsedDocument(
        text=text,
        title=path.stem,
        source_path=str(path),
        file_type="html",
    )


# helpers

def _clean_whitespace(text: str) -> str:
    """Normalise line endings, collapse excessive blank lines, strip edges."""
    import re
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{4,}", "\n\n\n", text)    # max 3 consecutive blank lines
    text = re.sub(r" {4,}", "   ", text)        # collapse long spaces
    return text.strip()


_CODE_EXTENSIONS = {
    ".py", ".js", ".ts", ".jsx", ".tsx", ".rs", ".go", ".java",
    ".c", ".cpp", ".h", ".cs", ".rb", ".php", ".sh", ".yaml",
    ".yml", ".toml", ".ini", ".env", ".sql", ".r", ".swift",
}

SUPPORTED_EXTENSIONS = {
    ".txt", ".md", ".rst", ".csv", ".json",
    ".pdf", ".docx", ".doc",
    ".html", ".htm",
} | _CODE_EXTENSIONS